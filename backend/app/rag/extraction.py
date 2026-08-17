"""Text extraction from uploaded documents.

This module is the one place that hands attacker-influenced bytes to a parsing library,
so it is deliberately narrow: it returns text and page numbers, nothing else. It never
evaluates a formula, never resolves an external reference, and never runs a macro —
DOCX/XLSX active content is simply not read (CLAUDE.md 4.2).

It runs inside a Celery worker, never in a request handler, so a file that manages to
hang or crash a parser takes down a replaceable worker instead of the API.
"""

from __future__ import annotations

import csv
import io
import os
import re
import tempfile
import zipfile
from collections.abc import Iterator
from dataclasses import dataclass
from pathlib import Path

from app.config import get_settings

#: Rows of a CSV grouped into one logical page, so citations can point somewhere useful.
CSV_ROWS_PER_PAGE = 200

#: Parts that identify what an OOXML container actually is, regardless of its extension.
_OOXML_MARKERS = {
    "docx": "word/document.xml",
    "xlsx": "xl/workbook.xml",
}

#: Zero-width and bidirectional control characters. Invisible in any viewer, but they land
#: inside tokens and split a word into two the embedding model has never seen. A PDF
#: extractor emits them routinely; they also happen to be the classic way to hide text in a
#: document, so removing them makes what is indexed match what a human reads.
#: Written as escapes rather than literals so the source stays greppable and reviewable.
_INVISIBLE = re.compile("[​-‏‪-‮⁠-⁤﻿]")

#: Any run of horizontal whitespace, including the non-breaking and typographic spaces PDF
#: extraction produces.
_HORIZONTAL_WS = re.compile(r"[^\S\n]+")

#: Three or more newlines — i.e. more than one blank line between blocks.
_EXTRA_BLANK_LINES = re.compile(r"\n{3,}")


def normalize_text(raw: str) -> str:
    """Collapse extraction artefacts into clean, embeddable text.

    Three transformations, each earning its place: invisible characters are dropped,
    runs of horizontal whitespace become a single space, and runs of blank lines collapse
    to one. A single blank line is preserved because it is the paragraph boundary the
    chunker splits on first — flattening it would cost the splitter its best separator.

    This runs before chunking, so chunk boundaries and the stored text agree. Normalizing
    afterwards would let a chunk be sized against whitespace that is then thrown away.
    """
    text = _INVISIBLE.sub("", raw.replace("\r\n", "\n").replace("\r", "\n"))
    text = _HORIZONTAL_WS.sub(" ", text)
    # Trailing spaces first, so a line of only spaces becomes empty and is then eligible
    # to be collapsed as a blank line rather than surviving as whitespace.
    text = "\n".join(line.strip() for line in text.split("\n"))
    return _EXTRA_BLANK_LINES.sub("\n\n", text).strip()


def count_words(text: str) -> int:
    """Whitespace-delimited word count.

    Deliberately naive: this is a size signal shown in the UI, not linguistic analysis,
    and a real tokenizer here would cost far more than the number is worth.
    """
    return len(text.split())


class ExtractionError(Exception):
    """A document could not be read. The message is recorded on the document row."""


@dataclass(frozen=True)
class ExtractedPage:
    """One unit of source text, with the locator a citation will quote."""

    #: 1-based page (PDF), sheet index (XLSX), or row block (CSV). None when not paginated.
    page: int | None
    text: str
    #: Human-readable locator for citations, e.g. "page 4" or "sheet 'Q3 Revenue'".
    label: str


def _guard_zip_bomb(path: Path) -> zipfile.ZipFile:
    """Open an OOXML container, refusing one that expands beyond the configured ceiling.

    The upload size cap bounds the *compressed* bytes only; a few hundred KB of ZIP can
    expand to gigabytes and exhaust the worker. The declared uncompressed sizes are
    checked first because reading the archive is what we are trying to avoid.
    """
    limit = get_settings().max_extracted_bytes
    try:
        archive = zipfile.ZipFile(path)
    except zipfile.BadZipFile as exc:
        raise ExtractionError("File is not a readable Office document.") from exc

    total = sum(info.file_size for info in archive.infolist())
    if total > limit:
        archive.close()
        raise ExtractionError(
            f"Document expands to {total // (1024 * 1024)} MB, over the "
            f"{limit // (1024 * 1024)} MB extraction limit."
        )
    return archive


def _assert_ooxml_kind(path: Path, kind: str) -> None:
    """Confirm the container really holds the document type its extension claims."""
    with _guard_zip_bomb(path) as archive:
        names = set(archive.namelist())
    marker = _OOXML_MARKERS[kind]
    if marker not in names:
        raise ExtractionError(f"File does not contain the parts of a valid .{kind} document.")


def _extract_pdf(path: Path) -> Iterator[ExtractedPage]:
    import pymupdf

    try:
        document = pymupdf.open(path)
    except Exception as exc:  # pymupdf raises bare Exception subclasses for damaged files
        raise ExtractionError("PDF could not be opened; it may be corrupt.") from exc

    with document:
        if document.needs_pass:
            raise ExtractionError("PDF is password-protected and cannot be indexed.")
        for index, page in enumerate(document, start=1):
            text = page.get_text("text").strip()
            if text:
                yield ExtractedPage(page=index, text=text, label=f"page {index}")


def _extract_docx(path: Path) -> Iterator[ExtractedPage]:
    import docx

    _assert_ooxml_kind(path, "docx")
    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise ExtractionError("DOCX could not be read; it may be corrupt.") from exc

    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    # Tables carry a lot of the meaning in policy documents, and python-docx keeps them
    # out of `paragraphs`, so they would silently vanish from the index otherwise.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if text:
        # DOCX has no reliable page boundaries without rendering it; the document is one unit.
        yield ExtractedPage(page=None, text=text, label="document")


def _extract_xlsx(path: Path) -> Iterator[ExtractedPage]:
    import openpyxl

    _assert_ooxml_kind(path, "xlsx")
    try:
        # data_only: read cached values, never evaluate a formula.
        # read_only: stream rows instead of building the whole workbook in memory.
        workbook = openpyxl.load_workbook(path, data_only=True, read_only=True)
    except Exception as exc:
        raise ExtractionError("XLSX could not be read; it may be corrupt.") from exc

    try:
        for index, sheet in enumerate(workbook.worksheets, start=1):
            lines = []
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if value is None else str(value).strip() for value in row]
                if any(cells):
                    lines.append(" | ".join(cells))
            if lines:
                yield ExtractedPage(
                    page=index,
                    text="\n".join(lines),
                    label=f"sheet '{sheet.title}'",
                )
    finally:
        workbook.close()


def _decode_text(path: Path) -> str:
    raw = path.read_bytes()
    # utf-8-sig drops a BOM if present; replacement keeps one bad byte from failing a
    # whole document, which matters for exports from legacy systems.
    return raw.decode("utf-8-sig", errors="replace")


def _extract_txt(path: Path) -> Iterator[ExtractedPage]:
    text = _decode_text(path).strip()
    if text:
        yield ExtractedPage(page=None, text=text, label="document")


def _extract_markdown(path: Path) -> Iterator[ExtractedPage]:
    """Markdown as plain text, deliberately unrendered.

    The syntax is left in place rather than stripped: `## Refund policy` embeds just as
    well as `Refund policy`, and a heading marker is a useful structural hint to the model
    reading a retrieved chunk. Nothing here parses links or HTML blocks, so no reference in
    the file is ever resolved or fetched — the same treatment .txt gets, which is what
    keeps a Markdown upload from being a request-forgery primitive.
    """
    text = _decode_text(path).strip()
    if text:
        yield ExtractedPage(page=None, text=text, label="document")


def _extract_csv(path: Path) -> Iterator[ExtractedPage]:
    content = _decode_text(path)
    try:
        dialect = csv.Sniffer().sniff(content[:8192])
    except csv.Error:
        dialect = csv.excel  # type: ignore[assignment]

    reader = csv.reader(io.StringIO(content), dialect)
    try:
        rows = [row for row in reader if any(cell.strip() for cell in row)]
    except csv.Error as exc:
        raise ExtractionError("CSV could not be parsed.") from exc

    if not rows:
        return

    header, *body = rows
    header_line = " | ".join(cell.strip() for cell in header)
    if not body:
        yield ExtractedPage(page=1, text=header_line, label="rows 1-1")
        return

    for block_index, start in enumerate(range(0, len(body), CSV_ROWS_PER_PAGE), start=1):
        block = body[start : start + CSV_ROWS_PER_PAGE]
        # The header is repeated in every block so a chunk is interpretable on its own,
        # which is what the retriever will hand to the model.
        lines = [header_line] + [" | ".join(cell.strip() for cell in row) for row in block]
        first_row, last_row = start + 1, start + len(block)
        yield ExtractedPage(
            page=block_index,
            text="\n".join(lines),
            label=f"rows {first_row}-{last_row}",
        )


_EXTRACTORS = {
    "pdf": _extract_pdf,
    "docx": _extract_docx,
    "xlsx": _extract_xlsx,
    "csv": _extract_csv,
    "txt": _extract_txt,
    "md": _extract_markdown,
}


def extract_pages(source: Path | bytes, *, extension: str) -> list[ExtractedPage]:
    """Extract and normalize text from a stored upload, dispatching on its extension.

    ``source`` may be a file path or raw bytes. When bytes are provided, a temporary
    file is written so the existing extractors (which expect a Path) continue to work.

    The extension comes from the allowlist check at upload time, never from the
    filename at this point — by here the file is named after a UUID.

    Normalization happens here rather than in each extractor so every format gets it, and
    a page left empty by it is dropped: a page of nothing but whitespace would otherwise
    become a chunk of nothing, take up a retrieval slot, and cite a page with no content.
    """
    extractor = _EXTRACTORS.get(extension)
    if extractor is None:
        raise ExtractionError(f"No extractor is registered for '{extension}' files.")

    tmp_path: Path | None = None
    try:
        if isinstance(source, bytes):
            if not source:
                raise ExtractionError("The uploaded file is empty.")
            # The extractors below expect a path, so bytes are staged in a temporary
            # file named with the canonical extension — never with a user-supplied
            # name, so traversal is impossible by construction.
            fd, raw_name = tempfile.mkstemp(suffix=f".{extension}")
            os.close(fd)
            tmp_path = Path(raw_name)
            tmp_path.write_bytes(source)
            path = tmp_path
        else:
            path = source

        if not path.exists():
            raise ExtractionError("Stored file is missing.")

        pages = [
            ExtractedPage(page=page.page, text=normalized, label=page.label)
            for page in extractor(path)
            if (normalized := normalize_text(page.text))
        ]
        if not pages:
            raise ExtractionError("No readable text was found in this document.")
        return pages
    finally:
        # A staged temporary file is ours to remove; a caller-supplied path is not.
        if tmp_path is not None:
            tmp_path.unlink(missing_ok=True)


__all__ = [
    "CSV_ROWS_PER_PAGE",
    "ExtractedPage",
    "ExtractionError",
    "count_words",
    "extract_pages",
    "normalize_text",
]
