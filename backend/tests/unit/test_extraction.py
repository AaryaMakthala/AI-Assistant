"""Extraction returns text plus an accurate locator, and refuses hostile files."""

from __future__ import annotations

import zipfile
from pathlib import Path

import pytest

from app.config import get_settings
from app.rag.extraction import CSV_ROWS_PER_PAGE, ExtractionError, extract_pages

pytestmark = pytest.mark.usefixtures("valid_env")


def _write_pdf(path: Path, pages: list[str]) -> Path:
    import pymupdf

    document = pymupdf.open()
    for body in pages:
        page = document.new_page()
        page.insert_text((72, 72), body)
    document.save(path)
    document.close()
    return path


def _write_docx(path: Path, paragraphs: list[str], table: list[list[str]] | None = None) -> Path:
    import docx

    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table:
        created = document.add_table(rows=len(table), cols=len(table[0]))
        for row_index, row in enumerate(table):
            for col_index, value in enumerate(row):
                created.cell(row_index, col_index).text = value
    document.save(path)
    return path


def _write_xlsx(path: Path, sheets: dict[str, list[list[object]]]) -> Path:
    import openpyxl

    workbook = openpyxl.Workbook()
    workbook.remove(workbook.active)
    for title, rows in sheets.items():
        sheet = workbook.create_sheet(title=title)
        for row in rows:
            sheet.append(row)
    workbook.save(path)
    return path


async def test_pdf_pages_are_numbered_from_one(tmp_path: Path) -> None:
    path = _write_pdf(tmp_path / "doc.pdf", ["Refund policy text", "Second page content"])

    pages = extract_pages(path, extension="pdf")

    assert [page.page for page in pages] == [1, 2]
    assert [page.label for page in pages] == ["page 1", "page 2"]
    assert "Refund policy" in pages[0].text
    assert "Second page" in pages[1].text


def test_pdf_skips_pages_with_no_text(tmp_path: Path) -> None:
    """A blank page must not become an empty chunk that pollutes retrieval."""
    path = _write_pdf(tmp_path / "gappy.pdf", ["First", "", "Third"])

    pages = extract_pages(path, extension="pdf")

    assert [page.page for page in pages] == [1, 3]


def test_docx_includes_table_cells(tmp_path: Path) -> None:
    """python-docx keeps tables out of `paragraphs`, so they need explicit handling."""
    path = _write_docx(
        tmp_path / "handbook.docx",
        ["Employees may claim expenses."],
        table=[["Category", "Limit"], ["Travel", "500"]],
    )

    pages = extract_pages(path, extension="docx")

    assert len(pages) == 1
    assert "Employees may claim expenses." in pages[0].text
    assert "Travel | 500" in pages[0].text


def test_xlsx_yields_one_page_per_sheet_labelled_by_name(tmp_path: Path) -> None:
    path = _write_xlsx(
        tmp_path / "book.xlsx",
        {"Q3 Revenue": [["Region", "Total"], ["EMEA", 120]], "Notes": [["ok"]]},
    )

    pages = extract_pages(path, extension="xlsx")

    assert [page.label for page in pages] == ["sheet 'Q3 Revenue'", "sheet 'Notes'"]
    assert "EMEA | 120" in pages[0].text


def test_csv_repeats_the_header_in_every_block(tmp_path: Path) -> None:
    """Each block is retrieved on its own, so it has to be readable on its own."""
    rows = ["name,amount"] + [f"customer{i},{i}" for i in range(CSV_ROWS_PER_PAGE + 10)]
    path = tmp_path / "orders.csv"
    path.write_text("\n".join(rows), encoding="utf-8")

    pages = extract_pages(path, extension="csv")

    assert len(pages) == 2
    assert all(page.text.startswith("name | amount") for page in pages)
    assert pages[0].label == f"rows 1-{CSV_ROWS_PER_PAGE}"
    assert pages[1].label == f"rows {CSV_ROWS_PER_PAGE + 1}-{CSV_ROWS_PER_PAGE + 10}"


def test_txt_is_decoded_with_bom_stripped(tmp_path: Path) -> None:
    path = tmp_path / "notes.txt"
    path.write_bytes("﻿Hello world".encode())

    pages = extract_pages(path, extension="txt")

    assert pages[0].text == "Hello world"
    assert pages[0].page is None


def test_document_with_no_text_is_an_error(tmp_path: Path) -> None:
    path = tmp_path / "blank.txt"
    path.write_text("   \n  ", encoding="utf-8")

    with pytest.raises(ExtractionError, match="No readable text"):
        extract_pages(path, extension="txt")


def test_missing_file_is_an_error(tmp_path: Path) -> None:
    with pytest.raises(ExtractionError, match="missing"):
        extract_pages(tmp_path / "absent.pdf", extension="pdf")


def test_unregistered_extension_is_an_error(tmp_path: Path) -> None:
    path = tmp_path / "thing.bin"
    path.write_text("x", encoding="utf-8")

    with pytest.raises(ExtractionError, match="No extractor"):
        extract_pages(path, extension="bin")


def test_xlsx_bytes_named_docx_are_rejected(tmp_path: Path) -> None:
    """Both are ZIPs, so the upload sniff cannot tell them apart — extraction must."""
    path = _write_xlsx(tmp_path / "sheet.xlsx", {"S": [["a"]]})
    disguised = tmp_path / "disguised.docx"
    disguised.write_bytes(path.read_bytes())

    with pytest.raises(ExtractionError, match="valid .docx"):
        extract_pages(disguised, extension="docx")


def test_corrupt_ooxml_is_rejected(tmp_path: Path) -> None:
    path = tmp_path / "broken.docx"
    path.write_bytes(b"PK\x03\x04 not really a zip")

    with pytest.raises(ExtractionError, match="not a readable Office document"):
        extract_pages(path, extension="docx")


def test_zip_bomb_is_refused_before_being_read(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A small archive that expands enormously must not be unpacked."""
    monkeypatch.setenv("MAX_EXTRACTED_BYTES", "1024")
    get_settings.cache_clear()

    path = tmp_path / "bomb.docx"
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", "A" * 100_000)

    with pytest.raises(ExtractionError, match="extraction limit"):
        extract_pages(path, extension="docx")

    get_settings.cache_clear()
